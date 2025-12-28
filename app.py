import streamlit as st
import google.generativeai as genai
from apify_client import ApifyClient
import urllib.parse
from datetime import datetime, date
from docx import Document
from io import BytesIO
import time
from PIL import Image, ImageDraw, ImageFont
import requests
from gtts import gTTS
from google.api_core.exceptions import ResourceExhausted, NotFound
from google.generativeai.types import HarmCategory, HarmBlockThreshold

# --- CONFIGURACIÓN VISUAL ---
st.set_page_config(page_title="SpyTool Pro: Ebook 2.0 Edition 📚", layout="wide")

# --- GESTIÓN DE SECRETOS ---
api_key_google = st.secrets.get("GOOGLE_API_KEY", None)
api_key_apify = st.secrets.get("APIFY_API_TOKEN", None)

# --- BLINDAJE ---
SAFETY_SETTINGS = {
    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
}

# --- FUNCIONES ---
def crear_word(titulo_libro, capitulos_guardados):
    doc = Document()
    doc.add_heading(titulo_libro, 0)
    for cap in capitulos_guardados:
        doc.add_heading(cap['titulo'], level=1)
        doc.add_paragraph(cap['contenido'])
        doc.add_page_break()
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def texto_a_audio(texto, idioma='es'):
    try:
        tts = gTTS(text=texto, lang=idioma, slow=False)
        buffer = BytesIO()
        tts.write_to_fp(buffer)
        buffer.seek(0)
        return buffer
    except: return None

def consultar_gemini_robusto(prompt, api_key, model_name_principal, lista_modelos_disponibles, stream=False):
    if not api_key: return None
    genai.configure(api_key=api_key)
    try:
        model = genai.GenerativeModel(model_name_principal)
        if stream: return model.generate_content(prompt, stream=True, safety_settings=SAFETY_SETTINGS)
        else: return model.generate_content(prompt, safety_settings=SAFETY_SETTINGS).text
    except (ResourceExhausted, NotFound):
        st.toast(f"⚠️ Motor {model_name_principal} ocupado. Buscando respaldo...", icon="🔄")
        time.sleep(1) 
        for backup_name in lista_modelos_disponibles:
            if backup_name != model_name_principal:
                try:
                    model_backup = genai.GenerativeModel(backup_name)
                    if stream: return model_backup.generate_content(prompt, stream=True, safety_settings=SAFETY_SETTINGS)
                    else: return model_backup.generate_content(prompt, safety_settings=SAFETY_SETTINGS).text
                except: continue
        st.error("🚦 Tráfico extremo. Espera 1 min.")
        return None
    except Exception as e:
        st.error(f"Error técnico: {e}")
        return None

# --- ESTADO ---
if 'borrador_libro' not in st.session_state: st.session_state['borrador_libro'] = []
if 'mis_modelos' not in st.session_state: st.session_state['mis_modelos'] = []

# --- BARRA LATERAL ---
with st.sidebar:
    st.header("⚙️ Configuración")
    if not api_key_google:
        api_key_google = st.text_input("1. Google API Key:", type="password")
        st.warning("⚠️ Clave no detectada en Secrets. Úsala manual.")
    else:
        st.success("✅ Google Key: Conectada")
        
    modelo_seleccionado = "models/gemini-1.5-flash" 
    if api_key_google:
        try:
            genai.configure(api_key=api_key_google)
            try:
                lista_reales = []
                for m in genai.list_models():
                    if 'generateContent' in m.supported_generation_methods:
                        lista_reales.append(m.name)
                st.session_state['mis_modelos'] = lista_reales
            except:
                lista_reales = ["models/gemini-1.5-flash", "models/gemini-pro"]
            if lista_reales:
                index_defecto = 0
                for i, nombre in enumerate(lista_reales):
                    if "gemini-1.5-flash" in nombre: index_defecto = i
                modelo_seleccionado = st.selectbox("🤖 Motor Principal:", lista_reales, index=index_defecto)
        except: pass
    
    st.divider()
    if not api_key_apify: api_key_apify = st.text_input("2. Apify Token (Opcional):", type="password")
    else: st.success("✅ Apify Token: Conectado")
    
    st.divider()
    if len(st.session_state['borrador_libro']) > 0:
        if st.button("🗑️ Reiniciar Libro"):
            st.session_state['borrador_libro'] = []
            st.rerun()

# --- INTERFAZ PRINCIPAL ---
st.title("📚 SpyTool Pro: Ebook 2.0 Factory")

if not api_key_google:
    st.info("👋 Configura tus llaves para empezar.")
    st.stop()

tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs(["📡 Radar", "🏭 Fábrica 2.0", "🎨 Portada Pro", "📢 Marketing", "🌐 Landing Page", "🎧 Extras", "🌪️ Embudo"])

# PESTAÑA 1: RADAR (CON BRASIL)
with tab1:
    st.header("Investigación de Mercado")
    modo = st.radio("Modo:", ["🤖 Automático", "✍️ Manual"], horizontal=True)
    if modo == "🤖 Automático":
        c1, c2 = st.columns(2)
        with c1: keyword = st.text_input("Nicho:", value="Productividad")
        with c2: pais = st.selectbox("País:", ["US", "ES", "MX", "BR"]) 
        if st.button("🚀 Buscar"):
            if not api_key_apify: st.error("Falta Token Apify.")
            else:
                try:
                    client = ApifyClient(api_key_apify)
                    encoded = urllib.parse.quote(keyword)
                    run_input = { "startUrls": [{ "url": f"https://www.facebook.com/ads/library/?active_status=active&ad_type=all&country={pais}&q={encoded}&search_type=keyword_unordered&media_type=all" }], "maxItems": 3 }
                    with st.spinner("Buscando..."):
                        run = client.actor("apify/facebook-ads-scraper").call(run_input=run_input, timeout_secs=60)
                    items = list(client.dataset(run["defaultDatasetId"]).iterate_items())
                    st.success(f"{len(items)} anuncios encontrados.")
                    for i, item in enumerate(items):
                        txt = item.get('adBody') or item.get('primaryText') or ""
                        with st.container(border=True):
                            st.text_area("Copy", txt, height=60, key=f"t{i}")
                            if st.button("🧠 Usar este Concepto", key=f"b{i}"):
                                st.session_state['tema'] = txt[:800]
                                st.success("Concepto cargado a la Fábrica.")
                except Exception as e: st.error(f"Error: {e}")
    else:
        tm = st.text_area("Pega tu idea o anuncio ganador:")
        if st.button("Analizar"):
            st.session_state['tema'] = tm[:800]
            st.success("Concepto cargado.")

# PESTAÑA 2: FÁBRICA 2.0 (ACTUALIZADA Y BLINDADA)
with tab2:
    st.header("🏭 Fábrica de Contenido 'Actionable'")
    tema = st.session_state.get('tema', 'Sin tema definido.')
    st.info(f"📌 Concepto Base: {tema[:100]}...")
    
    c_m, c_w = st.columns(2)
    with c_m:
        st.subheader("1. Estructura")
        if st.button("Generar Índice 'Tool-kit'"):
            prompt_indice = f"""
            Actúa como estratega de Info-productos Best-Seller.
            Crea un índice para un 'Workbook/Ebook Accionable' sobre: {tema}.
            NO quiero capítulos teóricos aburridos.
            Quiero títulos atractivos orientados a resultados (Ej: 'Reto Día 1', 'Tu Plan de Acción', 'La Técnica Secreta').
            Incluye secciones de 'Hoja de Trabajo' y 'Auditoría'.
            """
            res = consultar_gemini_robusto(prompt_indice, api_key_google, modelo_seleccionado, st.session_state['mis_modelos'])
            if res: st.session_state['mapa'] = res
        if 'mapa' in st.session_state: st.markdown(st.session_state['mapa'])
        
    with c_w:
        st.subheader("2. Redacción Inteligente")
        tit_cap = st.text_input("Título del Capítulo a escribir:")
        tipo_contenido = st.selectbox("Tipo de Contenido:", ["Lección + Ejercicio (Workbook)", "Checklist de Acción", "Reto de 24 Horas", "Híbrido (Texto + Video QR)"])
        inst = st.text_area("Detalles extra:")
        
        if st.button("✍️ Escribir Capítulo Pro"):
            # AQUI ESTA EL BLINDAJE DE CANTIDAD Y ESTILO
            prompt = f"""
            Actúa como redactor de Ebooks Best-Seller estilo 2026.
            Tema: {tema}.
            Capítulo: '{tit_cap}'.
            Formato Elegido: '{tipo_contenido}'.
            
            REGLAS DE ORO (Ebook 2.0 - LECTURA RÁPIDA):
            1. CONTROL DE LONGITUD: MÁXIMO 800 palabras. Sé extremadamente conciso. Si puedes decirlo en una frase, no uses un párrafo.
            2. ESTILO VISUAL: Usa muchas negritas, listas (bullets) y emojis. EVITA MUROS DE TEXTO.
            3. FORMATO WORKBOOK: Incluye obligatoriamente espacios para rellenar (usa lineas: __________).
            4. ACCIONABLE: Si es Checklist, usa casillas [ ].
            5. HÍBRIDO: Incluye un marcador visual que diga: > **[📱 ESCANEA AQUÍ EL QR PARA VER EL VIDEO EXPLICATIVO]**
            6. VALOR AGREGADO IA: Al final, incluye un recuadro con un "Prompt de ChatGPT" que el lector pueda usar.
            
            Instrucciones extra: {inst}.
            Salida en Markdown estético y limpio.
            """
            cont = st.empty()
            full = ""
            res = consultar_gemini_robusto(prompt, api_key_google, modelo_seleccionado, st.session_state['mis_modelos'], stream=True)
            if res:
                for ch in res:
                    try: 
                        if ch.text: full += ch.text; cont.markdown(full + "▌")
                    except: pass
                cont.markdown(full)
                st.session_state['temp_cap'] = full; st.session_state['temp_tit'] = tit_cap
                
        if 'temp_cap' in st.session_state:
            if st.button("💾 Guardar en Libro"):
                st.session_state['borrador_libro'].append({"titulo":st.session_state['temp_tit'], "contenido":st.session_state['temp_cap']})
                st.success("Capítulo guardado con éxito.")
                del st.session_state['temp_cap']
                st.rerun()
                
    if len(st.session_state['borrador_libro']) > 0:
        st.divider()
        st.download_button("📥 Descargar Ebook Completo (.docx)", crear_word("Mi Best-Seller", st.session_state['borrador_libro']), "Ebook_Accionable.docx")

# PESTAÑA 3: PORTADA PRO
with tab3:
    st.header("🎨 Portada de Alto Impacto")
    t_l = st.text_input("Título Libro:", placeholder="El Método 30 Días")
    st_l = st.selectbox("Estilo Visual:", ["Minimalista 'Apple'", "Bold Typography (Letras Gigantes)", "Estilo Revista Moderna", "3D Abstracto"])
    if st.button("🧠 Crear Prompt Ideogram"):
        with st.spinner("Diseñando concepto..."):
            prompt_base = f"Prompt para Ideogram AI. Ebook Cover design. Title: '{t_l}'. Style: {st_l}. Concept: '{tema}'. High contrast, professional, bestseller aesthetic. NO cluttered text. Big bold fonts. Clean layout."
            res = consultar_gemini_robusto(prompt_base, api_key_google, modelo_seleccionado, st.session_state['mis_modelos'])
            if res: st.code(res, language="text"); st.success("Copia esto en Ideogram.ai")

# PESTAÑA 4: MARKETING
with tab4:
    st.header("📢 Marketing & Ventas")
    tema_marketing = st.session_state.get('tema', 'Sin tema')
    tab_copy, tab_visual = st.tabs(["✍️ Anuncios", "🎬 Guiones Reels"])
    with tab_copy:
        if st.button("Generar Ads"):
            prompt = f"Escribe 3 Ads para Facebook sobre {tema_marketing}. Enfócate en el 'Dolor' y la 'Solución Rápida' (Low Ticket). Usa emojis y CTA claros."
            res = consultar_gemini_robusto(prompt, api_key_google, modelo_seleccionado, st.session_state['mis_modelos'])
            if res: st.markdown(res)
    with tab_visual:
        if st.button("Guiones Reels"):
            p = f"3 Guiones Reels 15s para vender '{tema_marketing}'. Estilo: Problema -> Agitación -> Solución (El Ebook). Formato Tabla. Sé directo."
            res = consultar_gemini_robusto(p, api_key_google, modelo_seleccionado, st.session_state['mis_modelos'])
            if res: st.markdown(res)

# PESTAÑA 5: LANDING PAGE
with tab5:
    st.header("🌐 Landing Page de Conversión")
    prod = st.text_input("Nombre Producto:", value="Kit de Acción 30 Días")
    prec = st.text_input("Precio Oferta:", value="$9 USD")
    if st.button("🏗️ Generar HTML"):
        prompt_web = f"HTML5 landing page moderna para '{prod}' ({tema}). Precio {prec}. Estilo 'Sales Letter' corta. Fondo blanco, letra negra legible. Botones rojos de compra. Responsive. Solo código HTML puro."
        res = consultar_gemini_robusto(prompt_web, api_key_google, modelo_seleccionado, st.session_state['mis_modelos'])
        if res:
            clean = res.replace("```html","").replace("```","")
            st.download_button("Descargar HTML", clean, "index.html")
            st.components.v1.html(clean, height=400, scrolling=True)

# PESTAÑA 6: EXTRAS
with tab6:
    st.header("🎧 Producción Audio & Legal")
    col_a, col_l = st.columns(2)
    with col_a:
        st.subheader("Audiobook (Order Bump)")
        if len(st.session_state['borrador_libro']) > 0:
            titulos = [c['titulo'] for c in st.session_state['borrador_libro']]
            sel = st.selectbox("Capítulo:", titulos)
            cont_cap = next((c['contenido'] for c in st.session_state['borrador_libro'] if c['titulo'] == sel), "")
            if st.button("Convertir a MP3"):
                with st.spinner("Grabando..."):
                    ab = texto_a_audio(cont_cap)
                    if ab: st.audio(ab); st.download_button("Descargar MP3", ab, f"{sel}.mp3")
    with col_l:
        st.subheader("Textos Legales")
        emp = st.text_input("Tu Marca:")
        mail = st.text_input("Tu Email:")
        if st.button("Generar"):
            p = f"Textos legales HTML (Privacidad, Descargo) para {emp} ({mail})."
            res = consultar_gemini_robusto(p, api_key_google, modelo_seleccionado, st.session_state['mis_modelos'])
            if res: st.markdown(res)

# PESTAÑA 7: EMBUDOS
with tab7:
    st.header("🌪️ Estrategia de Embudo")
    tema_funnel = st.session_state.get('tema', 'Sin tema')
    col_bump, col_upsell = st.columns(2)
    with col_bump:
        st.subheader("Order Bump ($7-$19)")
        tipo_bump = st.selectbox("Idea:", ["Audiobook", "Plantilla Notion", "Pack de Prompts IA"])
        if st.button("Crear Oferta Bump"):
            p = f"Texto corto persuasivo para Order Bump: {tipo_bump} relacionado con {tema_funnel}. Precio ridículo, valor alto. Usa formato Título + Beneficio."
            res = consultar_gemini_robusto(p, api_key_google, modelo_seleccionado, st.session_state['mis_modelos'])
            if res: st.info(res)
    with col_upsell:
        st.subheader("Upsell ($27+)")
        tipo_upsell = st.selectbox("Idea:", ["Masterclass Video", "Asesoría Grupal", "Comunidad VIP"])
        if st.button("Guion VSL Upsell"):
            p = f"Guion video ventas para Upsell: {tipo_upsell}. Cliente ya compró el ebook. Ahora véndele aceleración. Sé agresivo con la escasez."
            res = consultar_gemini_robusto(p, api_key_google, modelo_seleccionado, st.session_state['mis_modelos'])
            if res: st.markdown(res)
